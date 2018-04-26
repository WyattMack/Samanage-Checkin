import selenium, sys, datetime, os, requests, base64, json
from docx import Document
from selenium import webdriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import UnexpectedAlertPresentException
from selenium.common.exceptions import TimeoutException

#Saves username and password from a txt document in C:users\user\ Username and password must be exactly like this -> username:password
credentials = {}
with open("pass.txt", 'r') as passw:
    for line in passw:
        un, pw = line.strip().split(':')
        credentials[un]=pw
     
#assigns Firefox(geckodriver) as the browser to use and specifies the browser path
driver = webdriver.Firefox(executable_path=r'C:\Geckodriver\geckodriver.exe')
driver.implicitly_wait(5)

#Goes to assetexplorer:8080 to login
driver.get('http://assetexplorer:8080')

#fills login info then signs in
username = driver.find_element_by_css_selector('#username')
username.send_keys(un)
password = driver.find_element_by_css_selector('#password')
password.send_keys(pw)
logOn = driver.find_element_by_css_selector('#domainListSelect > select')
logOn.send_keys('INFUSIONSOFT.PHX')
submit = driver.find_element_by_css_selector('#logindetailstable > tbody > tr:nth-child(5) > td:nth-child(2) > table > tbody > tr > td.pleft6.bor-left-fff > a > input')
submit.click()

"""
try:
    WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.ID, 'searchText')))
except TimeoutException:
        print("Timed out")
        driver.quit()
"""

search = driver.find_element_by_css_selector('#searchText')
search.send_keys(sys.argv[1])

submit = driver.find_element_by_css_selector('.button')
submit.click()

select1 = driver.find_element_by_css_selector('#AssetView_TABLE > tbody > tr:nth-child(4) > td:nth-child(3) > a')
select1.click()

assetElement = driver.find_elements_by_xpath("/html/body/table/tbody/tr[1]/td/table/tbody/tr[2]/td/table/tbody/tr/td[2]/table/tbody/tr[2]/td/div/div/table/tbody/tr[4]/td/table[1]/tbody/tr[2]/td/table/tbody/tr/td[1]/table/tbody/tr[2]/td[3]")
AssetTag = [x.text for x in assetElement]

LastUserElement = driver.find_elements_by_xpath("/html/body/table/tbody/tr[1]/td/table/tbody/tr[2]/td/table/tbody/tr/td[2]/table/tbody/tr[2]/td/div/div/table/tbody/tr[4]/td/table[1]/tbody/tr[1]/td/table/tbody/tr[2]/td/table/tbody/tr/td[1]/table/tbody/tr[2]/td[3]")
LastUser = [x.text for x in LastUserElement]

ComputerTypeElement = driver.find_elements_by_xpath("/html/body/table/tbody/tr[1]/td/table/tbody/tr[2]/td/table/tbody/tr/td[2]/table/tbody/tr[2]/td/div/div/table/tbody/tr[4]/td/table[1]/tbody/tr[1]/td/table/tbody/tr[2]/td/table/tbody/tr/td[2]/table/tbody/tr[1]/td[3]")
ComputerType = [x.text for x in ComputerTypeElement]

ExpiresElement = driver.find_elements_by_xpath("/html/body/table/tbody/tr[1]/td/table/tbody/tr[2]/td/table/tbody/tr/td[2]/table/tbody/tr[2]/td/div/div/table/tbody/tr[4]/td/table[1]/tbody/tr[2]/td/table/tbody/tr/td[2]/table/tbody/tr[5]/td[3]")
Expires = [x.text for x in ExpiresElement]

DiskElement = driver.find_elements_by_xpath('//*[@id="CIDet_Diskspace"]')
Disk = [x.text for x in DiskElement]

RAMElement = driver.find_elements_by_xpath('//*[@id="CIDet_TotalMemory"]')
RAM = [x.text for x in RAMElement]

SerialElement = driver.find_elements_by_xpath('/html/body/table/tbody/tr[1]/td/table/tbody/tr[2]/td/table/tbody/tr/td[2]/table/tbody/tr[2]/td/div/div/table/tbody/tr[4]/td/table[1]/tbody/tr[2]/td/table/tbody/tr/td[1]/table/tbody/tr[3]/td[3]')
Serial = [x.text for x in SerialElement]

ServiceTagElement = driver.find_elements_by_xpath('//*[@id="CIDet_ServiceTag"]')
ServiceTag = [x.text for x in ServiceTagElement]

ciName = driver.find_element_by_xpath('//*[@id="CIDet_CIName"]').text
computerName = ciName.split(".")[0]

edit = driver.find_element_by_css_selector('#startListMenuItems > tbody:nth-child(1) > tr:nth-child(1) > td:nth-child(5) > a:nth-child(1)')
edit.click()

getUser= driver.find_element_by_xpath('//*[@id="s2id_user"]').text

menu = driver.find_element_by_css_selector('#select2-chosen-1')
menu.click()
 
searchAsset = driver.find_element_by_css_selector('#s2id_autogen1_search')
searchAsset.send_keys('In Store')
searchAsset.send_keys(u'\ue007')

save = driver.find_element_by_css_selector('input.btn:nth-child(3)')
save.click()

now = datetime.datetime.now()

#opens a word document and saves all the asset info then prints it

document = Document()

table = document.add_table(rows=3, cols=4)
cell0 = table.cell(0, 0)
cell1 = table.cell(0, 1)
cell2 = table.cell(0, 2)
cell3 = table.cell(0, 3)

r1cell0 = table.cell(1, 0)
r1cell1 = table.cell(1, 1)
r1cell2 = table.cell(1, 2)
r1cell3 = table.cell(1, 3)

r2cell0 = table.cell(2, 0)
r2cell1 = table.cell(2, 1)
r2cell2 = table.cell(2, 2)
r2cell3 = table.cell(2, 3)


row0 = table.rows[0]
cell0 = row0.cells[0]
cell0.text = 'Asset Tag:'
cell1 = row0.cells[1]
cell1.text = AssetTag
cell2 = row0.cells[2]
cell2.text = 'Model:'
cell3 = row0.cells[3]
cell3.text = ComputerType


row1 = table.rows[1]
cell4 = row1.cells[0]
cell4.text = 'Specs:'
cell5 = row1.cells[1]
cell5.text = 'Ram: '+RAM[0] #str(RAM).strip+'HDD: '+str(Disk).strip
cell6 = row1.cells[2]
cell6.text = 'LAU:'
cell7 = row1.cells[3]
if len(getUser) < 1:
    getUser = computerName
cell7.text = getUser

row2 = table.rows[2]
cell8 = row2.cells[0]
cell8.text = 'SN/Service Tag:'
cell9 = row2.cells[1]
cell9.text = Serial[0]# + '/' + str(ServiceTag).strip
cell10 = row2.cells[2]
cell10.text = 'In Store:'
cell11 = row2.cells[3]
cell11.text = now.strftime('%Y-%m-%d' + ', ' + Expires[0])
#if Expires[0] == [-]
#    Expires[0] = 

run0 = cell0.paragraphs[0].runs[0]
#run1 = cell1.paragraphs[0].runs[0]
run2 = cell2.paragraphs[0].runs[0]
#run3 = cell3.paragraphs[0].runs[0]
run4 = cell4.paragraphs[0].runs[0]
#run5 = cell5.paragraphs[0].runs[0]
run6 = cell6.paragraphs[0].runs[0]
#run7 = cell7.paragraphs[0].runs[0]
run8 = cell8.paragraphs[0].runs[0]
#run9 = cell9.paragraphs[0].runs[0]
run10 = cell10.paragraphs[0].runs[0]

run0.font.bold = True
#run1.font.bold = True
run2.font.bold = True
#run3.font.bold = True
run4.font.bold = True
#run5.font.bold = True
run6.font.bold = True
#run7.font.bold = True
run8.font.bold = True
#run9.font.bold = True
run10.font.bold = True


document.save('AssetInfo.docx')

os.startfile("AssetInfo.docx", "print")

#trys to remove it from Jamf and prints success message or prints nothing to delete 
try:

    #Saves username and password from a txt document in C:users\user\ Username and password must be exactly like this -> username:password
    credentials = {}
    with open("jamf.txt", 'r') as passw:
        for line in passw:
            UsernameVar, PasswordVar = line.strip().split(':')
            credentials[UsernameVar]=PasswordVar

    # base URL of JSS
    jssUrl = 'https://infusionsoft.jamfcloud.com/JSSResource'
    computerUrl = jssUrl + '/computers/name/' + computerName

    response = requests.get(computerUrl, headers = {'Accept': 'application/json'}, auth=(UsernameVar, PasswordVar))

    macID = (response_json['computer']['general']['id'])

    MacIdUrl = jssUrl + '/computercommands/command/DeviceLock/passcode/996688/id/' + str(macID)

    IDresponse = requests.post(MacIdUrl, headers = {'Accept': 'application/json'}, auth=(UsernameVar, PasswordVar))

    print (IDresponse)
    

#    response = requests.delete(computerUrl, headers = {'Accept': 'application/json'}, auth=(UsernameVar, PasswordVar))

#    print ('Successfully removed from Jamf')
    
except:
    print ("Nothing to delete")



