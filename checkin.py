import sys, datetime, os, requests, base64, json, selenium
import tkinter as Tk
from tkinter import simpledialog
from docx import Document
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import UnexpectedAlertPresentException
from selenium.webdriver.common.action_chains import ActionChains
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.by import By
from selenium.webdriver.support.select import Select

#Collect search information
query = Tk.Tk()
query.withdraw()
input = simpledialog.askstring(title='Asset/SN/or Computer Name',
                                    prompt='Search Criteria:\nPlease Use Asset Tag\nOr Computer Name:')

#Saves username and password from a txt document in C:users\user\ Username and password must be exactly like this -> username:password
credentials = {}
os.chdir(r'C:\\Users\\username\\')
with open("pass.txt", 'r') as passw:
    for line in passw:
        un, pw = line.strip().split(':')
        credentials[un]=pw
     
#assigns Firefox(geckodriver) as the browser to use and specifies the browser path
driver = webdriver.Firefox(executable_path=r'C:\Geckodriver\geckodriver.exe')
driver.implicitly_wait(10)

#Goes to infusionsoft.samanage.com to login
driver.get('https://infusionsoft.samanage.com/')

#fills login info then signs in
username = driver.find_element_by_css_selector('#login')
username.send_keys(un)
password = driver.find_element_by_css_selector('#password')
password.send_keys(pw)
submit = driver.find_element_by_css_selector('#form_button')
submit.click()

#Begins searching for the computer info input
search = driver.find_element_by_css_selector('.SearchInput__searchIconStyle___3QJL2')
search.click()

submit = driver.find_element_by_css_selector('.SearchInput__input___29PgA')
submit.send_keys(input)
submit.send_keys(Keys.ENTER)

#Clicks into the computer record if found, and prompts to search again if no record is found
try:
    record = driver.find_element_by_css_selector('.object_link')
    record.click()

except: 
    'selenium.common.exceptions.NoSuchElementException: Message: Unable to locate element: .object_link'
    input = simpledialog.askstring(title='Asset/SN/or Computer Name',
                prompt='Search Criteria:\nPlease Use Asset Tag\nOr Computer Name:')
    submit = driver.find_element_by_css_selector('.SearchInput__input___29PgA').clear()
    submit = driver.find_element_by_css_selector('.SearchInput__input___29PgA')
    submit.send_keys(input)
    submit.send_keys(Keys.ENTER)

#Scrape machine name
machine = driver.find_elements_by_xpath('/html/body/div[3]/div/div[2]/div[2]/div/div/div[2]/div[2]/div[2]/div[1]')
machine = [x.text for x in machine]
machine = '1'.join(machine)
machine = machine.split('\n')[1]
print(machine)

#Scrape last active user information
lau = driver.find_elements_by_xpath('/html/body/div[3]/div/div[2]/div[2]/div/div/div[2]/div[2]/div[2]/div[2]')
lau = [x.text for x in lau]
lau = '1'.join(lau)
lau = lau.split('\n')[1]
print(lau)

#Scrape asset tag
tag = driver.find_elements_by_xpath('/html/body/div[3]/div/div[2]/div[2]/div/div/div[2]/div[4]/div[1]/div[1]/div[1]/span[1]')
tag = [x.text for x in tag]
tag = '1'.join(tag)
print(tag)

#Get current date and time
t = datetime.datetime.now()

#Add a note with check-in date, last active user, and machine name for historical recording.
#Define note by assigning a driver path
note = driver.find_element_by_css_selector('div.softwareDetailsArea:nth-child(4) > div:nth-child(1) > div:nth-child(1) > div:nth-child(2) > span:nth-child(1)')

#Have driver scroll down the page
driver.execute_script("window.scrollTo(0, 600);")

#Assign button as driver path where button will appear
button = driver.find_element_by_xpath('/html/body/div[3]/div/div[2]/div[2]/div/div/div[2]/div[10]/div/div/div[1]/div[1]/div/span[2]')

#Have hover make the button visible
hover = ActionChains(driver).move_to_element(button)
hover.perform()

#Click and send the note info, then save
note = driver.find_element_by_xpath('/html/body/div[3]/div/div[2]/div[2]/div/div/div[2]/div[10]/div/div/div[1]/div[1]/div/span[2]')
note.click()
note = driver.find_element_by_xpath('//*[@id="hardware__98706"]')
note.click()
note.send_keys(Keys.END)
note.send_keys('Checked in ' + t.strftime('%x') + ' -- LAU: ' + lau + ' ' + machine + '\n')
note = driver.find_element_by_xpath('/html/body/div[3]/div/div[2]/div[2]/div/div/div[2]/div[10]/div/div/div[1]/div[1]/div/a[1]')
note.click()
note.send_keys(Keys.HOME)

#Click into "Hardware" tab
hw = driver.find_element_by_xpath('/html/body/div[3]/div/div[2]/div[2]/div/div/div[1]/div/strong/ul/li[5]/a')
hw.click()

#Scrape computer model
model = driver.find_elements_by_xpath('/html/body/div[3]/div/div[2]/div[2]/div/div/div[2]/div[3]/div[2]/div[2]')
model = [x.text for x in model]
model = '1'.join(model)
model = model.split('\n')[1]
print(model)

#Scrape serial number
sn = driver.find_elements_by_xpath('/html/body/div[3]/div/div[2]/div[2]/div/div/div[2]/div[3]/div[2]/div[3]/div[1]/span[1]')
sn = [x.text for x in sn]
sn = '1'.join(sn)
print(sn)

#Scrape RAM
RAM = driver.find_elements_by_xpath('/html/body/div[3]/div/div[2]/div[2]/div/div/div[2]/div[3]/div[1]/div[2]')
RAM = [x.text for x in RAM]
RAM = '1'.join(RAM)
RAM = RAM.split('\n')[1]
print(RAM)

#Click into "Lifecycle" tab
lc = driver.find_element_by_xpath('/html/body/div[3]/div/div[2]/div[2]/div/div/div[1]/div/strong/ul/li[6]/a')
lc.click()

#Scrape warranty info
wnty = driver.find_elements_by_xpath('/html/body/div[3]/div/div[2]/div[2]/div/div/div[2]/div[1]/div[2]/table/tbody/tr[2]/td[5]')
wnty = [x.text for x in wnty]
wnty = '1'.join(wnty)
print(wnty)

#Opens a word document and saves all the asset info then prints it

document = Document()

table = document.add_table(rows=3, cols=4)
cell0 = table.cell(0, 0)
cell1 = table.cell(0, 1)
cell2 = table.cell(0, 2)
cell3 = table.cell(0, 3)

r1cell0 = table.cell(1, 0)
r1cell1 = table.cell(1, 1)
r1cell2 = table.cell(1, 2)
r1cell3 = table.cell(1, 3)

r2cell0 = table.cell(2, 0)
r2cell1 = table.cell(2, 1)
r2cell2 = table.cell(2, 2)
r2cell3 = table.cell(2, 3)


row0 = table.rows[0]
cell0 = row0.cells[0]
cell0.text = 'Asset Tag:'
cell1 = row0.cells[1]
cell1.text = tag
cell2 = row0.cells[2]
cell2.text = 'Model:'
cell3 = row0.cells[3]
cell3.text = model


row1 = table.rows[1]
cell4 = row1.cells[0]
cell4.text = 'RAM:'
cell5 = row1.cells[1]
cell5.text = RAM[0:2] + ' GB'
cell6 = row1.cells[2]
cell6.text = 'LAU:'
cell7 = row1.cells[3]
#if len(getUser) < 1:
#    getUser = computerName
cell7.text = lau

row2 = table.rows[2]
cell8 = row2.cells[0]
cell8.text = 'S/N:'
cell9 = row2.cells[1]
cell9.text = sn
cell10 = row2.cells[2]
cell10.text = 'In Store:'
cell11 = row2.cells[3]
cell11.text = t.strftime('%x' + ' ( warranty expires: ' + wnty + ')')
#if Expires[0] == [-]
#    Expires[0] = 

run0 = cell0.paragraphs[0].runs[0]
#run1 = cell1.paragraphs[0].runs[0]
run2 = cell2.paragraphs[0].runs[0]
#run3 = cell3.paragraphs[0].runs[0]
run4 = cell4.paragraphs[0].runs[0]
#run5 = cell5.paragraphs[0].runs[0]
run6 = cell6.paragraphs[0].runs[0]
#run7 = cell7.paragraphs[0].runs[0]
run8 = cell8.paragraphs[0].runs[0]
#run9 = cell9.paragraphs[0].runs[0]
run10 = cell10.paragraphs[0].runs[0]

run0.font.bold = True
#run1.font.bold = True
run2.font.bold = True
#run3.font.bold = True
run4.font.bold = True
#run5.font.bold = True
run6.font.bold = True
#run7.font.bold = True
run8.font.bold = True
#run9.font.bold = True
run10.font.bold = True


document.save('AssetInfo.docx')

os.startfile('AssetInfo.docx')
os.startfile('AssetInfo.docx', 'print')